"""
app/extractors/ocr_extractor.py
─────────────────────────────────
Local OCR extractor — Tesseract (via pytesseract) or PaddleOCR.

CURRENT STATE: Functional stub.
Returns None for all fields with error_code='NOT_ENABLED'.
Real implementation is gated behind OCR_ENABLED=true in .env.

HOW TO ENABLE (real implementation):
1. Install Tesseract: `brew install tesseract`
2. Add `pytesseract` + `Pillow` + `pdf2image` to requirements.txt
3. Set OCR_ENABLED=true in .env
4. In `_run_ocr()`:
   a. If file_type == 'pdf': convert pages to images with pdf2image.
   b. Run pytesseract.image_to_string() on each page image.
   c. Concatenate full text.
5. Pass full text to `_extract_fields()` which uses regex/keyword heuristics
   to pull out V1 fields.

ALTERNATIVE: PaddleOCR (better accuracy for complex layouts):
   Replace pytesseract calls with PaddleOCR().ocr(img) and parse results.

WHY OCR is lower priority than Textract in tie-breaking:
Local OCR does not understand document structure (forms, tables),
so it has lower precision on key-value extraction from contracts.
"""

import os
import re
import uuid
from typing import Optional

from app.core.logging import get_logger
from app.extractors.base import V1_FIELDS, BaseExtractor, ExtractionResult

# ── A4: Junk-value filter ─────────────────────────────────────────────────────
# Section-header words that regex can wrongly capture as field values.
# Applied after every regex match before assigning.
_JUNK_VALUES = re.compile(
    r"^(Details|for\s+Notices?|Notices?|below\.?|above\.?|herein|parties|applicable|"
    r"following|hereinafter|the\s+agreement|this\s+agreement|schedule|"
    r"annex|appendix|exhibit|section|clause|article|page|date|signature|"
    r"as\s+follows|set\s+out\s+below)$",
    re.IGNORECASE,
)

# ── Regex patterns for V1 field extraction from raw OCR text ──────────────────
# These are applied after the full document text is extracted.

# A1: person_name — colon REQUIRED (not optional) to avoid matching section headers
# like "Party Details". Removed "name" and "party" — both appear in headings.
_RE_PERSON_NAME = re.compile(
    r"(?:signed\s+by|contractor|consultant|employee|vendor|representative|undersigned)\s*:\s*"
    r"([A-Z][A-Za-z .'-]{2,40})",
    re.IGNORECASE,
)

# company_name: colon required for the same reason.
_RE_COMPANY_NAME = re.compile(
    r"(?:company|entity|client|employer|organization|firm|service\s+provider)\s*:\s*"
    r"([A-Z][A-Za-z0-9 .,&'-]{2,60})",
    re.IGNORECASE,
)

# contract_date: ISO date, US date, or written date.
_RE_DATE = re.compile(
    r"\b(\d{4}-\d{2}-\d{2}"                           # YYYY-MM-DD
    r"|\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}"              # MM/DD/YYYY or MM-DD-YYYY
    r"|(?:January|February|March|April|May|June|July|August|September|October|November|December)"
    r"\s+\d{1,2}(?:st|nd|rd|th)?,?\s+\d{4})\b",       # Month DD, YYYY
    re.IGNORECASE,
)

# A3: contract_value — added INR/₹/Rs. (Indian contracts) + \n?\s* allows the
# amount to appear on the next line after the label keyword.
_RE_VALUE = re.compile(
    r"(?:value|amount|consideration|total|price|fee|compensation|payment)\s*[:\-]?\s*\n?\s*"
    r"((?:INR|₹|Rs\.?|USD|\$|£|€|GBP|EUR)\s*[\d,]+(?:\.\d{1,2})?"
    r"|[\d,]+(?:\.\d{1,2})?\s*(?:INR|USD|GBP|EUR)"
    r"|\$[\d,]+(?:\.\d{1,2})?)",
    re.IGNORECASE,
)

# Fallback: any INR/₹/Rs. or bare $ amount (two capture groups).
_RE_VALUE_FALLBACK = re.compile(
    r"(?:INR|₹|Rs\.?)\s*([\d,]+(?:\.\d{1,2})?)"
    r"|\$([\d,]+(?:\.\d{1,2})?)",
    re.IGNORECASE,
)

# A2: address — captured value must start with a digit.
# Real addresses start with numbers ("123 Main St", "5th Floor", "11th Floor").
# This blocks "for Notices", "hereinafter", etc.
_RE_ADDRESS = re.compile(
    r"(?:address|located at|principal place|place of business|address\s+for\s+notices?)\s*[:\-]\s*"
    r"(\d[^\n]{5,120})",
    re.IGNORECASE,
)

# Fix 1: address label and value on separate lines (common in Indian contracts).
# Matches the digit-starting address value on the line immediately after the label.
_RE_ADDRESS_NEXT_LINE = re.compile(
    r"(?:address|located at|principal place|place of business|address\s+for\s+notices?)"
    r"\s*[:\-]?\s*\n\s*(\d[^\n]{5,120})",
    re.IGNORECASE,
)

# Fix 3: extract person name from signature block cells.
# Matches "Authorized Signatory: Rahul Khanna" or "Authorised Signatory: Meera Iyer".
_RE_SIGNATORY = re.compile(
    r"(?:Authorized|Authorised)\s+Signatory\s*:\s*([A-Z][A-Za-z .'-]{2,40})",
    re.IGNORECASE,
)

# contract_date (labelled) — prefer dates next to "effective date", "dated as of", etc.
# These are far more reliable than the first date found anywhere in the document.
_RE_EFFECTIVE_DATE = re.compile(
    r"(?:effective\s+(?:as\s+of\s+)?date|dated\s+(?:as\s+of|this)\b|entered\s+into\s+(?:as\s+of|this)\b|agreement\s+date)\s*[:\-]?\s*"
    r"(\d{4}-\d{2}-\d{2}"
    r"|\d{1,2}[/\-]\d{1,2}[/\-]\d{2,4}"
    r"|(?:January|February|March|April|May|June|July|August|September|October|November|December)"
    r"\s+\d{1,2}(?:st|nd|rd|th)?,?\s+\d{4})",
    re.IGNORECASE,
)

logger = get_logger(__name__)

# ── B: KV label maps for table row parsing ────────────────────────────────────
# pdfplumber extracts table cells as "Label | Value" rows.
# These sets map normalised left-cell text to the target V1 field.
_KV_PERSON_LABELS = frozenset({
    "contractor", "consultant", "employee", "vendor", "representative",
    "signed by", "undersigned", "individual", "service provider",
    "name of contractor", "name of consultant", "name of employee",
    # Fix 3 + Fix 4: signature block labels + common variants
    "authorized signatory", "authorised signatory", "signatory",
    "authorized representative", "authorised representative",
})
_KV_COMPANY_LABELS = frozenset({
    "company", "client", "employer", "entity", "firm", "organization",
    "service provider", "company name", "client name", "employer name",
    "name of company", "name of client",
    # Fix 4: multi-party contract labels
    "party a", "party b", "party 1", "party 2", "first party", "second party",
    "purchaser", "buyer", "seller", "licensor", "licensee",
    "service recipient", "customer",
})
_KV_DATE_LABELS = frozenset({
    "effective date", "agreement date", "contract date", "date",
    "dated", "date of agreement", "commencement date", "start date",
    "execution date", "signing date",
})
_KV_VALUE_LABELS = frozenset({
    "fee", "amount", "total fee", "total amount", "contract value",
    "consideration", "total consideration", "payment", "total payment",
    "price", "contract price", "compensation", "total compensation", "value",
    "contract amount", "total contract value",
})
_KV_ADDRESS_LABELS = frozenset({
    "address", "registered address", "corporate address", "office address",
    "principal place of business", "place of business", "located at",
    "billing address", "mailing address",
    # Fix 1 + Fix 4: additional address label variants
    "address for notices", "address for notice", "notice address",
    "registered office", "registered office address", "principal office",
    "for notices", "correspondence address",
})


def _extract_from_kv_rows(kv_rows: list[str]) -> dict[str, Optional[str]]:
    """
    Parse pdfplumber table rows ("cell1 | cell2 | ...") as key-value pairs.

    The first non-empty cell is the key, the second is the value.
    Returns only fields that were found (missing fields absent, not None).
    KV results take priority over regex in _run_ocr().
    """
    result: dict[str, Optional[str]] = {}
    for row in kv_rows:
        parts = [p.strip() for p in row.split("|") if p.strip()]
        if len(parts) < 2:
            continue
        raw_key = parts[0].lower().strip()
        raw_val = parts[1].strip()
        if not raw_val or _JUNK_VALUES.match(raw_val):
            continue
        if raw_key in _KV_PERSON_LABELS and "person_name" not in result:
            # Fix 3: multi-line cell may be "Company\nAddress\nAuthorized Signatory: Name"
            # If signatory pattern found, that is the person name.
            sig_match = _RE_SIGNATORY.search(raw_val)
            result["person_name"] = sig_match.group(1).strip() if sig_match else raw_val.split("\n")[0].strip()
        elif raw_key in _KV_COMPANY_LABELS and "company_name" not in result:
            # Fix 3: take only first line of multi-line signature-block cell
            first_line = raw_val.split("\n")[0].strip()
            if not _JUNK_VALUES.match(first_line):
                result["company_name"] = first_line
            # Also opportunistically extract person from signatory in same cell
            if "person_name" not in result:
                sig_match = _RE_SIGNATORY.search(raw_val)
                if sig_match:
                    result["person_name"] = sig_match.group(1).strip()
        elif raw_key in _KV_DATE_LABELS and "contract_date" not in result:
            result["contract_date"] = raw_val
        elif raw_key in _KV_VALUE_LABELS and "contract_value" not in result:
            result["contract_value"] = raw_val
        elif raw_key in _KV_ADDRESS_LABELS and "address" not in result:
            # Take first line to avoid multi-line address + signatory collapse
            result["address"] = raw_val.split("\n")[0].strip()
    return result


# ── spaCy NER — lazy loader + helpers ────────────────────────────────────────
# The model is loaded once on first use and cached at module level.
# If spaCy or the model is unavailable the extractor silently falls back to regex.
_nlp = None  # None = not yet attempted; False = attempted and failed


def _load_nlp():
    """Return the cached spaCy nlp pipeline, loading it on first call."""
    global _nlp
    if _nlp is None:
        try:
            import spacy  # noqa: PLC0415
            _nlp = spacy.load("en_core_web_sm")
            logger.info("spaCy en_core_web_sm loaded for NER extraction")
        except Exception as exc:
            logger.warning(
                f"spaCy model unavailable — falling back to regex for name fields: {exc}"
            )
            _nlp = False  # sentinel: tried, failed, don't retry
    return _nlp if _nlp is not False else None


# Keywords whose proximity to a NER entity signals it is the target entity.
_PERSON_CONTEXT = re.compile(
    r"\b(?:contractor|employee|party|signed by|signature|consultant|vendor|"
    r"individual|freelancer|representative|undersigned)\b",
    re.IGNORECASE,
)
_ORG_CONTEXT = re.compile(
    r"\b(?:company|employer|client|entity|firm|corporation|inc|ltd|llc|"
    r"organization|provider|service\s+provider|hereinafter)\b",
    re.IGNORECASE,
)


def _ner_extract(text: str, nlp) -> tuple[Optional[str], Optional[str]]:
    """
    Run spaCy NER once and return (person_name, company_name).

    Selection strategy:
    - person_name: PERSON entity closest to a context keyword; fall back to the
      last PERSON entity in the doc (signature-block heuristic).
    - company_name: ORG entity closest to a context keyword; fall back to the
      first ORG entity (preamble heuristic — company is introduced early).

    Text is capped at 50 000 chars to keep spaCy fast on large contracts.
    """
    # Cap input — spaCy slows significantly beyond ~100k chars
    doc = nlp(text[:50_000])
    persons = [ent for ent in doc.ents if ent.label_ == "PERSON"]
    orgs = [ent for ent in doc.ents if ent.label_ == "ORG"]

    # Rejects spaCy ORG entities that look like address fragments
    # (e.g. "5th Floor", "3rd Floor", "Suite 100").
    _RE_ADDR_FRAGMENT = re.compile(
        r"^\d+(?:st|nd|rd|th)\s+(?:floor|suite|block|tower|building|wing)\b"
        r"|^(?:floor|suite|block|tower|building|wing)\s+\d+",
        re.IGNORECASE,
    )

    def _clean_ent(ent_text: str) -> str:
        """Take first line only; strip leading non-capital chars (e.g. 'g Party A ...')."""
        first_line = ent_text.split("\n")[0].strip()
        # Fix 2: strip any leading chars before the first uppercase letter
        # (spaCy entity boundary sometimes starts mid-word)
        return re.sub(r'^[^A-Z]+', '', first_line)

    # ── person_name ───────────────────────────────────────────────────────────
    person_name: Optional[str] = None
    if persons:
        kw_positions = [m.start() for m in _PERSON_CONTEXT.finditer(text)]
        if kw_positions:
            person_name = _clean_ent(
                min(
                    persons,
                    key=lambda e: min(abs(e.start_char - p) for p in kw_positions),
                ).text
            )
        else:
            # No context keywords found — signature block is at the end.
            person_name = _clean_ent(persons[-1].text)

    # ── company_name ──────────────────────────────────────────────────────────
    company_name: Optional[str] = None
    if orgs:
        kw_positions = [m.start() for m in _ORG_CONTEXT.finditer(text)]
        if kw_positions:
            company_name = _clean_ent(
                min(
                    orgs,
                    key=lambda e: min(abs(e.start_char - p) for p in kw_positions),
                ).text
            )
        else:
            # No context keywords — company is usually named in the preamble.
            company_name = _clean_ent(orgs[0].text)

    # Discard address-fragment false positives (e.g. "5th Floor")
    if company_name and _RE_ADDR_FRAGMENT.match(company_name):
        company_name = None

    return person_name, company_name


# Read OCR_ENABLED inside extract() at call time to avoid module-level import
# order issues with .env loading.


def _extract_fields_from_text(text: str) -> dict[str, Optional[str]]:
    """
    Extract V1 fields from raw text.

    person_name + company_name: spaCy NER first, regex fallback.
    contract_date / contract_value / address: regex (Phase 1 already handles these well).
    """
    result: dict[str, Optional[str]] = {f: None for f in V1_FIELDS}

    # ── person_name + company_name via spaCy NER ──────────────────────────────
    nlp = _load_nlp()
    if nlp:
        ner_person, ner_company = _ner_extract(text, nlp)
        result["person_name"] = ner_person
        result["company_name"] = ner_company

    # Regex fallback (A4: junk filter applied after every match).
    if not result["person_name"]:
        m = _RE_PERSON_NAME.search(text)
        if m and not _JUNK_VALUES.match(m.group(1).strip()):
            result["person_name"] = m.group(1).strip()

    if not result["company_name"]:
        m = _RE_COMPANY_NAME.search(text)
        if m and not _JUNK_VALUES.match(m.group(1).strip()):
            result["company_name"] = m.group(1).strip()

    # contract_date — prefer dates near effective-date labels; fall back to first date
    m = _RE_EFFECTIVE_DATE.search(text)
    if m:
        result["contract_date"] = m.group(1).strip()
    else:
        m = _RE_DATE.search(text)
        if m:
            result["contract_date"] = m.group(1).strip()

    # contract_value — prefer labelled match, then any currency-prefixed amount
    m = _RE_VALUE.search(text)
    if m:
        result["contract_value"] = m.group(1).strip()
    else:
        m = _RE_VALUE_FALLBACK.search(text)
        if m:
            # Two capture groups: group(1)=INR/₹/Rs, group(2)=$
            result["contract_value"] = (m.group(1) or m.group(2)).strip()

    # address — same-line first; if no match try value on next line after label (Fix 1)
    m = _RE_ADDRESS.search(text) or _RE_ADDRESS_NEXT_LINE.search(text)
    if m:
        result["address"] = m.group(1).split("\n")[0].strip()

    # Fix 3: signatory regex as final fallback for person_name
    if not result["person_name"]:
        m = _RE_SIGNATORY.search(text)
        if m:
            result["person_name"] = m.group(1).strip()

    return result


class OCRExtractor(BaseExtractor):
    """
    Extracts V1 fields from document text using local OCR (Tesseract/PaddleOCR).

    Priority in consolidation tie-breaking: LOWEST (after Textract and LLMs).
    """

    method = "ocr"

    def extract(
        self,
        document_id: uuid.UUID,
        file_path: str,
        file_type: str,
    ) -> list[ExtractionResult]:
        if not os.getenv("OCR_ENABLED", "false").lower() == "true":
            logger.info(f"OCR skipped for document_id={document_id} (OCR_ENABLED=false)")
            return self._null_results(document_id, error_code="NOT_ENABLED",
                                      error_message="Set OCR_ENABLED=true to activate")

        return self._run_ocr(document_id, file_path, file_type)

    def _run_ocr(
        self,
        document_id: uuid.UUID,
        file_path: str,
        file_type: str,
    ) -> list[ExtractionResult]:
        """
        Extract fields from a document in three layers (highest precision first):
          1. KV table rows  — structured "Label | Value" cells from pdfplumber
          2. spaCy NER      — PERSON / ORG entities for names (regex fallback)
          3. Regex patterns — dates, values, addresses fill any remaining gaps
        Then normalise contract_date → YYYY-MM-DD and contract_value → bare number.
        """
        # ── Step 1: Extract text + structured KV rows ─────────────────────────
        extraction = self._extract_text(document_id, file_path, file_type)
        if extraction is None:
            return self._null_results(
                document_id,
                error_code="TEXT_EXTRACTION_FAILED",
                error_message="Could not extract text from document",
            )
        raw_text, kv_rows = extraction

        # ── Step 2a: KV table rows first (highest precision) ──────────────────
        extracted: dict[str, Optional[str]] = {f: None for f in V1_FIELDS}
        extracted.update(_extract_from_kv_rows(kv_rows))

        # ── Step 2b: NER + regex fill any fields KV missed ────────────────────
        regex_result = _extract_fields_from_text(raw_text)
        for field in V1_FIELDS:
            if extracted[field] is None:
                extracted[field] = regex_result[field]

        # ── Step 3: Normalize → YYYY-MM-DD date, bare numeric value ──────────
        if extracted.get("contract_date"):
            try:
                from dateutil import parser as _du
                extracted["contract_date"] = _du.parse(
                    extracted["contract_date"], dayfirst=False
                ).strftime("%Y-%m-%d")
            except Exception:
                pass

        if extracted.get("contract_value"):
            import re as _re
            cleaned = _re.sub(
                r"[$£€₹]|INR|USD|GBP|EUR|Rs\.?", "",
                extracted["contract_value"], flags=_re.IGNORECASE,
            ).replace(",", "").strip()
            if cleaned:
                extracted["contract_value"] = cleaned

        logger.info(
            f"OCR extracted document_id={document_id} "
            f"fields_found={sum(1 for v in extracted.values() if v)}/{len(V1_FIELDS)} "
            f"kv_rows={len(kv_rows)} text_length={len(raw_text)}"
        )

        return [
            ExtractionResult(
                document_id=document_id,
                method=self.method,
                field=f,
                value=extracted.get(f),
                evidence_snippet=raw_text[:200] if extracted.get(f) else None,
            )
            for f in V1_FIELDS
        ]

    def _extract_text(
        self,
        document_id: uuid.UUID,
        file_path: str,
        file_type: str,
    ) -> Optional[tuple[str, list[str]]]:
        """
        Returns (full_text, kv_rows).

        kv_rows are table rows kept as separate "cell1 | cell2" strings so
        _extract_from_kv_rows() can parse them as structured key-value pairs
        rather than running regex on flattened text.

        PDFs: pdfplumber first (digital text layer). Tesseract fallback for
        scanned/image PDFs only (kv_rows=[] in that case).
        DOCX: python-docx paragraphs + tables.
        """
        try:
            if file_type == "pdf":
                # ── Attempt 1: native text layer via pdfplumber ───────────────
                try:
                    import pdfplumber  # lazy import
                    para_parts: list[str] = []
                    kv_rows: list[str] = []
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text() or ""
                            if page_text:
                                para_parts.append(page_text)
                            for table in page.extract_tables() or []:
                                for row in table:
                                    cells = [
                                        str(c or "").strip() for c in row
                                        if c and str(c).strip()
                                    ]
                                    if len(cells) >= 2:
                                        # Keep as KV row — label + value still split
                                        kv_rows.append(" | ".join(cells))
                                    elif len(cells) == 1:
                                        para_parts.append(cells[0])
                    full_text = "\n".join(para_parts + kv_rows).strip()
                    if full_text:
                        logger.info(
                            f"OCR: pdfplumber extracted document_id={document_id} "
                            f"kv_rows={len(kv_rows)}"
                        )
                        return full_text, kv_rows
                except Exception:
                    pass  # pdfplumber unavailable or failed — fall through to Tesseract

                # ── Attempt 2: Tesseract (scanned / image PDFs) ──────────────
                logger.info(
                    f"OCR: pdfplumber returned no text, falling back to Tesseract "
                    f"for document_id={document_id}"
                )
                from pdf2image import convert_from_path  # lazy import
                import pytesseract                        # lazy import
                pages = convert_from_path(file_path, dpi=200)
                return "\n".join(pytesseract.image_to_string(page) for page in pages), []

            elif file_type == "docx":
                from docx import Document as DocxDocument  # lazy import
                doc = DocxDocument(file_path)
                para_parts = [p.text for p in doc.paragraphs if p.text.strip()]
                kv_rows = []
                for table in doc.tables:
                    for row in table.rows:
                        cells = [
                            cell.text.strip() for cell in row.cells
                            if cell.text.strip()
                        ]
                        if len(cells) >= 2:
                            kv_rows.append(" | ".join(cells))
                        elif len(cells) == 1:
                            para_parts.append(cells[0])
                full_text = "\n".join(para_parts + kv_rows)
                return (full_text, kv_rows) if full_text else None

            else:
                logger.warning(
                    f"OCR: unsupported file_type={file_type!r} for document_id={document_id}"
                )
                return None

        except Exception as exc:
            logger.error(
                f"OCR text extraction failed document_id={document_id} "
                f"file_type={file_type}: {exc}"
            )
            return None

    def _null_results(
        self,
        document_id: uuid.UUID,
        error_code: Optional[str] = None,
        error_message: Optional[str] = None,
    ) -> list[ExtractionResult]:
        return [
            ExtractionResult(
                document_id=document_id,
                method=self.method,
                field=f,
                value=None,
                error_code=error_code,
                error_message=error_message,
            )
            for f in V1_FIELDS
        ]

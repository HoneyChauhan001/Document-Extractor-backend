"""
app/extractors/gpt_extractor.py
─────────────────────────────────
OpenAI GPT extractor — LLM-based field extraction with strict JSON output.

CURRENT STATE: Functional stub.
Real calls are gated behind OPENAI_API_KEY being set in .env.

HOW IT WORKS (when enabled):
1. Extract raw text from the document (PDF→text via pdfplumber, DOCX via python-docx).
2. Send a structured prompt to GPT-4o (or gpt-4-turbo) that forces JSON output
   with exactly the 5 V1 fields.
3. Parse the JSON response deterministically.
4. Any JSON parse failure → all fields null + error recorded.

PROMPT CONTRACT (per llm-prompting-json-schema SKILL):
- Response must be a JSON object with EXACTLY these keys:
  person_name, company_name, contract_date, contract_value, address
- No extra keys allowed.
- null if not found; do not guess or infer.
- contract_date must be YYYY-MM-DD if parseable, else null.
- contract_value must be the numeric string with currency stripped (e.g. "50000").

HOW TO ENABLE:
1. pip install openai pdfplumber python-docx
2. Add secrets to .env: OPENAI_API_KEY=sk-...  OPENAI_MODEL=gpt-4o
3. Set GPT_ENABLED=true in .env
4. Replace stub body with real openai.chat.completions.create() call.
"""

import json
import os
import re
import uuid
from typing import Optional

from app.core.logging import get_logger
from app.extractors.base import V1_FIELDS, BaseExtractor, ExtractionResult

logger = get_logger(__name__)

# Read GPT_ENABLED + model inside extract() at call time — not at module level.
_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")  # model is safe to read at module level

# The strict system prompt — instructs GPT to return only the 5 V1 fields as JSON.
_SYSTEM_PROMPT = """You are a contract data extraction assistant.
Extract the following fields from the contract text provided.
Return a JSON object with EXACTLY these keys and no others:
  person_name, company_name, contract_date, contract_value, address

Rules:
- Use null (JSON null) if the field is not present. Do NOT guess or infer.
- contract_date: return as YYYY-MM-DD. Parse any date format (e.g. "March 15, 2024" → "2024-03-15", "15/03/2024" → "2024-03-15"). If not parseable, return null.
- contract_value: return the numeric amount only as a string with no currency symbols, letters, or commas (e.g. "$50,000.00" → "50000.00", "USD 1,200" → "1200"). Return null if not found.
- person_name: the primary signing individual's full name. Signature blocks near the END of the document are the most reliable source — prioritise them.
- company_name: the legal entity name of the company party (e.g. "Acme Corp", "TechCo Inc."). If two companies are named, prefer the one that is not the individual's employer.
- address: the primary business or registered address in the contract.
- Do NOT include any explanation, markdown, or keys beyond the five listed above.

Example:
Contract snippet: "This Service Agreement is entered into as of March 15, 2024 between Acme Corp, located at 123 Main St, New York, NY 10001 (\"Company\") and John Smith (\"Contractor\"). Total compensation: $50,000.00. /s/ John Smith"
Expected JSON output:
{\"person_name\": \"John Smith\", \"company_name\": \"Acme Corp\", \"contract_date\": \"2024-03-15\", \"contract_value\": \"50000.00\", \"address\": \"123 Main St, New York, NY 10001\"}"""


class GPTExtractor(BaseExtractor):
    """
    Extracts V1 fields by prompting OpenAI GPT with a strict JSON-schema prompt.

    Priority in consolidation: LLM (higher than OCR, lower than Textract in tie-breaks).
    """

    method = "gpt"

    def extract(
        self,
        document_id: uuid.UUID,
        file_path: str,
        file_type: str,
    ) -> list[ExtractionResult]:
        if not os.getenv("GPT_ENABLED", "false").lower() == "true":
            logger.info(f"GPT skipped for document_id={document_id} (GPT_ENABLED=false)")
            return self._null_results(document_id, error_code="NOT_ENABLED",
                                      error_message="Set GPT_ENABLED=true and provide OPENAI_API_KEY")

        # ── Read document text ─────────────────────────────────────────────────
        doc_text = _read_document_text(file_path, file_type)
        if not doc_text:
            return self._null_results(document_id, error_code="EMPTY_DOCUMENT",
                                      error_message="Could not extract text from document")

        # ── Call GPT ───────────────────────────────────────────────────────────
        try:
            import openai  # imported lazily — only needed when enabled
            client = openai.OpenAI()  # reads OPENAI_API_KEY from environment
            response = client.chat.completions.create(
                model=_MODEL,
                response_format={"type": "json_object"},  # forces JSON output
                messages=[
                    {"role": "system", "content": _SYSTEM_PROMPT},
                    # Smart chunk: first 4k + last 4k captures preamble (date, company,
                    # address) AND signature block / financials at the end.
                    {"role": "user", "content": f"Contract text:\n\n{_smart_chunk(doc_text)}"},
                ],
                temperature=0,      # deterministic output — no creative guessing
                timeout=30,         # 30s timeout; fail fast rather than hang
            )
            raw_json = response.choices[0].message.content or "{}"
        except Exception as exc:
            logger.error(f"GPT API call failed for document_id={document_id}: {exc}")
            return self._null_results(document_id, error_code="API_ERROR",
                                      error_message=str(exc))

        # ── Parse + post-process ───────────────────────────────────────────────
        results = _parse_llm_json(raw_json, document_id, self.method)
        return _post_process_results(results)

    def _null_results(
        self,
        document_id: uuid.UUID,
        error_code: Optional[str] = None,
        error_message: Optional[str] = None,
    ) -> list[ExtractionResult]:
        return [
            ExtractionResult(
                document_id=document_id,
                method=self.method,
                field=f,
                value=None,
                error_code=error_code,
                error_message=error_message,
            )
            for f in V1_FIELDS
        ]


# ── Shared helpers (used by both GPT and Claude extractors) ──────────────────

def _smart_chunk(text: str, first_n: int = 4000, last_n: int = 4000) -> str:
    """
    Return a context window that covers both the document preamble and the end.

    WHY: Contracts put the date/company/address in the opening clause but hide
    person_name and contract_value near the signature block at the end.
    Sending only the first N chars misses those fields entirely.
    If the document fits in first_n + last_n chars it is returned unchanged.
    """
    if len(text) <= first_n + last_n:
        return text
    return (
        text[:first_n]
        + "\n\n[... middle section omitted for brevity ...]\n\n"
        + text[-last_n:]
    )


def _read_document_text(file_path: str, file_type: str) -> Optional[str]:
    """
    Extract plain text + table content from a PDF or DOCX file.

    Tables are converted to pipe-separated rows so the LLM can read
    key-value pairs like "Contract Value | $50,000" that live in tables.
    """
    try:
        if file_type == "pdf":
            import pdfplumber
            parts: list[str] = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text:
                        parts.append(page_text)
                    # Tables — extract and flatten to readable rows.
                    for table in page.extract_tables() or []:
                        for row in table:
                            row_text = " | ".join(
                                str(cell or "").strip() for cell in row if cell
                            )
                            if row_text.strip():
                                parts.append(row_text)
            return "\n".join(parts) or None
        elif file_type == "docx":
            from docx import Document
            doc = Document(file_path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            # Tables — iterate rows and cells explicitly.
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts) or None
    except Exception as exc:
        logger.error(f"Text extraction failed for {file_path}: {exc}")
    return None


def _normalize_date(raw: str) -> Optional[str]:
    """Parse any date string and normalize to YYYY-MM-DD. Returns raw value on failure."""
    try:
        from dateutil import parser as _du
        return _du.parse(raw, dayfirst=False).strftime("%Y-%m-%d")
    except Exception:
        return raw


def _normalize_currency(raw: str) -> Optional[str]:
    """Strip currency symbols/codes and commas; return bare numeric string."""
    cleaned = re.sub(r"[$£€]|USD|GBP|EUR", "", raw, flags=re.IGNORECASE)
    cleaned = cleaned.replace(",", "").strip()
    return cleaned if cleaned else None


def _post_process_results(results: list[ExtractionResult]) -> list[ExtractionResult]:
    """
    Normalize contract_date → YYYY-MM-DD and contract_value → bare numeric string.

    Applied after LLM JSON parsing so that even when the model ignores the
    formatting rules in the prompt (common with smaller models), we still get
    clean values suitable for cross-method comparison in consolidation.
    """
    for result in results:
        if result.value is None:
            continue
        if result.field == "contract_date":
            result.value = _normalize_date(result.value)
        elif result.field == "contract_value":
            result.value = _normalize_currency(result.value)
    return results


def _parse_llm_json(
    raw_json: str,
    document_id: uuid.UUID,
    method: str,
) -> list[ExtractionResult]:
    """
    Parse the LLM's JSON response into ExtractionResult list.

    Strict contract:
    - Only the 5 V1 field keys are read; all others are ignored.
    - Invalid/non-parseable JSON → all nulls + PARSE_FAILURE error.
    - Values that are not strings or None are coerced to strings.

    This function is shared by both GPTExtractor and ClaudeExtractor.
    """
    try:
        parsed = json.loads(raw_json)
        if not isinstance(parsed, dict):
            raise ValueError("LLM response is not a JSON object")
    except (json.JSONDecodeError, ValueError) as exc:
        logger.warning(f"LLM JSON parse failure for document_id={document_id} "
                       f"method={method}: {exc} | raw={raw_json[:200]}")
        return [
            ExtractionResult(
                document_id=document_id,
                method=method,
                field=f,
                value=None,
                error_code="PARSE_FAILURE",
                error_message=f"JSON parse error: {exc}",
            )
            for f in V1_FIELDS
        ]

    # Build results in canonical V1_FIELDS order — always.
    results = []
    for f in V1_FIELDS:
        raw_val = parsed.get(f)  # None if key missing
        # Coerce to str; treat JSON null, empty string, "null" string as None.
        if raw_val is None or raw_val == "" or str(raw_val).lower() == "null":
            value = None
        else:
            value = str(raw_val).strip() or None
        results.append(
            ExtractionResult(
                document_id=document_id,
                method=method,
                field=f,
                value=value,
            )
        )
    return results
